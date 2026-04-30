import streamlit as st
import os, io, json, base64, datetime, requests, ssl, codecs
import pandas as pd
import numpy as np
import cv2
import plotly.express as px
from PIL import Image, ImageFilter
from docx import Document
import google.genai as genai
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# --- [0] 설정 파일 로드 ---
def load_config():
    with open('config.json', 'r', encoding='utf-8') as f:
        return json.load(f)

cfg = load_config()

# --- [1] 페이지 설정 ---
st.set_page_config(
    page_title=f"{cfg['institution']['abbr']} {cfg['institution']['app_title']}",
    layout="wide",
    page_icon="🚨"
)

# --- [2] 변수 및 구글 ID (동적 로드) ---
DRIVE_FOLDER_ID = cfg['google_api']['drive_folder_id']
SPREADSHEET_ID = cfg['google_api']['spreadsheet_id']
MODEL_ID = "gemini-flash-latest"

# --- [추가] 실제 구글 드라이브에 파일을 업로드하는 함수 ---
def upload_to_drive(file_name, file_content, mime_type):
    """
    구글 드라이브의 특정 폴더로 파일을 업로드하고, 
    공유 가능한 링크(webViewLink)를 반환합니다.
    """
    if drive_service is None:
        st.error("구글 드라이브 서비스가 연결되지 않았습니다.")
        return None
    
    try:
        # 1. 파일 메타데이터 설정
        file_metadata = {
            'name': file_name,
            'parents': [DRIVE_FOLDER_ID]
        }
        
        # 2. 파일 콘텐츠 준비
        media = MediaIoBaseUpload(
            io.BytesIO(file_content), 
            mimetype=mime_type, 
            resumable=True
        )
        
        # 3. 드라이브에 파일 생성 (fields에 webViewLink를 요청하는 것이 핵심!)
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, webViewLink', # 👈 여기를 수정해야 시트에 링크가 찍힙니다.
            supportsAllDrives=True,
            supportsTeamDrives=True
        ).execute()
        
        file_id = file.get('id')
        
        # 4. 파일 권한 설정 (이 코드가 있어야 워크스페이스 계정에서 사진이 바로 보입니다)
        user_permission = {
            'type': 'anyone',
            'role': 'reader',
        }
        drive_service.permissions().create(
            fileId=file_id,
            body=user_permission,
            supportsAllDrives=True,
            supportsTeamDrives=True
        ).execute()
        
        # 5. 생성된 파일의 공유 링크 반환
        return file.get('webViewLink')
        
    except Exception as e:
        st.error(f"구글 드라이브 업로드 중 에러 발생: {e}")
        return None

if "gcp_service_account" in st.secrets:
    try:
        creds_info = st.secrets["gcp_service_account"]
        
        # private_key 내의 실제 줄바꿈 문자 처리 (가장 흔한 오류 원인)
        if isinstance(creds_info, (dict, st.runtime.secrets.AttrDict)):
            creds_dict = dict(creds_info)
            creds_dict["private_key"] = creds_dict["private_key"].replace("\\n", "\n")
            
            SCOPES = ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/spreadsheets']
            creds = service_account.Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
            
            drive_service = build('drive', 'v3', credentials=creds)
            sheets_service = build('sheets', 'v4', credentials=creds)
    except Exception as e:
        st.error(f"⚠️ 인증 설정 오류: {e}")

else:
    st.error("Secrets 설정에서 'gcp_service_account'를 찾을 수 없습니다.")

# (이후 기존의 CSS 설정 및 나머지 코드를 이어 붙이시면 됩니다.)
# 주의: 아래쪽에 있는 st.set_page_config(page_title="KYWA AI 위험성평가 시스템", ...) 코드는 삭제하세요.

# [수정 2] 파라미터 이름을 unsafe_allow_html=True 로 변경
st.markdown("""
    <style>
    /* 1. 상단 헤더 메뉴 및 푸터 제거 (추가됨) */
    header {visibility: hidden;}
    footer {visibility: hidden;}
    #MainMenu {visibility: hidden;}
    
    /* 2. 상단 여백 조절 (추가됨) */
    .block-container {
        padding-top: 0rem;
        padding-bottom: 1rem;
    }

    /* 기존 코드 내용 유지 */
    html, body, [data-testid="stWidgetLabel"] p {
        color: var(--text-color);
    }
    
    .stDataFrame {
        width: 100% !important;
    }
    
    img {
        max-width: 100%;
        filter: brightness(var(--image-brightness, 1));
    }
    </style>
    """, unsafe_allow_html=True)

# 1. 환경 설정 및 보안 우회 (필요한 경우)
os.environ['PYTHONHTTPSVERIFY'] = '0'
ssl._create_default_https_context = ssl._create_unverified_context

# 2. 페이지 설정 및 세션 초기화
st.set_page_config(page_title="KYWA AI 위험성평가 시스템", layout="wide", page_icon="🚨")

if "analysis_results" not in st.session_state:
    st.session_state.analysis_results = None
if "final_data" not in st.session_state:
    st.session_state.final_data = None

# 3. 모델 및 클라이언트 설정 (최신 google-genai 방식)
try:
    if "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"]
        
        # 클라이언트 객체 생성 (기존 genai.configure 대체)
        # 2026년 기준, 별도의 transport 설정 없이도 최적화된 통신을 지원합니다.
        client = genai.Client(api_key=api_key)
        
        # 모델 이름 정의 (2026년 표준인 gemini-2.0-flash 권장, gemini-flash-latest 사용중임)
        # 만약 기존 모델을 유지하고 싶다면 'gemini-1.5-flash' 등을 입력하세요.
        model_name = "model=MODEL_ID" 
        
    else:
        st.error("Secrets에 'GEMINI_API_KEY'가 설정되지 않았습니다.")
        st.stop()
except Exception as e:
    st.error(f"API 설정 오류가 발생했습니다: {e}")
    st.stop()

# --- 도구 함수 (Word/Excel 생성) ---
def create_docx(data):
    doc = Document()
    doc.add_heading('KYWA AI 위험성평가 결과 보고서', 0)
    for item in data:
        doc.add_paragraph(f"분류: {item.get('category')}")
        doc.add_paragraph(f"장소: {item.get('location')}")
        doc.add_paragraph(f"상황: {item.get('scenario')}")
        doc.add_paragraph(f"등급: {item.get('grade')} (점수: {item.get('score')})")
        doc.add_paragraph(f"대책: {item.get('solution')}")
        doc.add_paragraph("-" * 20)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_excel(data):
    df = pd.DataFrame(data)
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')
    return bio.getvalue()



# --- [2단계] 구글 드라이브/시트 전송 함수 추가 ---

def upload_photo_to_drive(file_obj, filename):
    try:
        # Apps Script 웹 앱 URL (방금 복사한 주소)
        apps_script_url = "https://script.google.com/macros/s/AKfycbwMhipDH9zMVajhbD2LBXGgnJdaqs3oHmatjqtvAXWL0PXhInk6tqsqRcb6MJkZFChm/exec"
        
        file_obj.seek(0)
        img_base64 = base64.b64encode(file_obj.getvalue()).decode('utf-8')
        
        payload = {
            "filename": filename,
            "fileBase64": img_base64
        }
        
        # POST 요청으로 사진 전송
        response = requests.post(apps_script_url, json=payload)
        res_data = response.json()
        
        return res_data.get("url", "업로드 실패")
    except Exception as e:
        return f"업로드 실패: {str(e)}"

def append_row_to_sheet(row_data):
    try:
        # [확인 필요] 실제 시트 탭 이름과 공백이 정확히 일치해야 합니다.
        # 만약 구글폼 연동 시트라면 보통 '설문지 응답 1' 입니다.
        range_name = "'설문지 응답 시트1'!A1" 
        
        body = {'values': [row_data]}
        sheets_service.spreadsheets().values().append(
            spreadsheetId=SPREADSHEET_ID, 
            range=range_name,
            valueInputOption='USER_ENTERED', 
            insertDataOption='INSERT_ROWS', # [추가] 새 행을 삽입하며 추가하도록 명시
            body=body
        ).execute()
        return True
    except Exception as e:
        st.error(f"시트 저장 실패: {str(e)}")
        return False

# --- [추가] 버튼 스타일 설정 ---
st.markdown("""
    <style>
    /* 모든 Streamlit 버튼 스타일 수정 */
    div.stButton > button {
        background-color: #ff4b4b !important; /* 기본 붉은색 */
        color: white !important;
        border: none !important;
        padding: 0.5rem 1rem !important;
        border-radius: 0.5rem !important;
        font-weight: bold !important;
        transition: all 0.3s ease !important;
    }

    /* 마우스 호버(Hover) 시 효과 */
    div.stButton > button:hover {
        background-color: #ff3333 !important; /* 마우스 올렸을 때 더 진한 빨강 */
        color: white !important;
        border: none !important;
        transform: scale(1.01); /* 아주 살짝 커지는 효과 */
    }
    
    /* Word/Excel 저장 버튼 등 일반 버튼도 동일 적용을 원치 않으시면 위 범위를 좁힐 수 있습니다 */
    </style>
""", unsafe_allow_html=True)

# --- 4. 스타일 및 헤더 디자인 (안전 모드) ---
st.markdown("""
    <style>
    /* 버튼 스타일 */
    div.stButton > button {
        background-color: #ff4b4b !important;
        color: white !important;
        font-weight: bold !important;
        border-radius: 0.5rem !important;
        transition: all 0.3s ease !important;
    }
    div.stButton > button:hover {
        background-color: #ff3333 !important;
        transform: scale(1.01);
    }
    /* 로고 및 타이틀 스타일 */
    .logo-img { cursor: pointer; display: block; margin-top: 2px; }
    .refresh-title { text-decoration: none !important; color: inherit !important; cursor: pointer; }
    .refresh-title:hover { color: #FF4B4B !important; }
    </style>
""", unsafe_allow_html=True)

# 변수 초기화
local_logo_url = None

# 경로를 최대한 보수적으로 탐색 (에러 발생 시 앱이 멈추지 않도록 try-except 처리)
try:
    # 1. 현재 파일의 위치 파악
    current_path = os.path.dirname(os.path.abspath(__file__))
    logo_file = os.path.join(current_path, "kywa_logo.png")
    
    # 2. 파일이 실제로 존재할 때만 읽기 시도
    if os.path.exists(logo_file):
        with open(logo_file, "rb") as f:
            data = f.read()
            if data:
                encoded = base64.b64encode(data).decode()
                local_logo_url = f"data:image/png;base64,{encoded}"
except Exception as e:
    # 서버 로그에만 에러를 남기고 앱은 계속 실행됨
    print(f"Logo loading error: {e}")

# --- [3] 동적 헤더: 로고 및 타이틀 (교체 구간 시작) ---
h_col1, h_col2 = st.columns([1, 4])

with h_col1:
    # config.json의 경로를 사용하여 로고를 불러옵니다.
    try:
        inst_logo_base64 = base64.b64encode(open(cfg['institution']['logo_path'], "rb").read()).decode()
        st.markdown(f'''
            <a href="{cfg['institution']['website_url']}" target="_blank">
                <img src="data:image/png;base64,{inst_logo_base64}" width="250" class="logo-img">
            </a>
        ''', unsafe_allow_html=True)
    except Exception:
        # 로고 파일이 없을 경우 텍스트로 대체
        st.markdown(f"### {cfg['institution']['abbr']}")

with h_col2:
    st.markdown(f"""
        <a href="/" target="_self" style="text-decoration:none;">
            <h1 style='margin-bottom: 0;'>🚨 {cfg['institution']['name']} {cfg['institution']['app_title']}</h1>
        </a>
        <p style='color: gray; margin-top: 0;'>{cfg['institution']['app_subtitle']}</p>
    """, unsafe_allow_html=True)

st.divider()

# --- [4] 입력 섹션 (시설명 및 부서명 동적 로드) ---
col1, col2 = st.columns(2)

with col1:
    st.markdown("### **🏢 점검 대상 정보**")
    
    # config.json의 facilities 리스트를 사용합니다.
    시설명_list = cfg['ui_options']['facilities']
    selected_facility = st.radio("• 시설명 선택 (필수)", 시설명_list, horizontal=True)
    
    # config.json의 departments 리스트를 사용합니다.
    담당부서_list = cfg['ui_options']['departments']
    selected_dept = st.selectbox("• 담당 부서 선택 (필수)", 담당부서_list)
    
    st.markdown("### **📝 현장 상황 설명**")
    placeholder_text = "<예 시>\n1. ㅇㅇ시설물 파손 및 노후화\n2. 통로 적치물로 인한 전도 위험\n(자세히 작성할수록 정확한 결과가 나옵니다.)"
    user_description = st.text_area("• 상황 설명 입력 (권장)", placeholder=placeholder_text, height=150)

with col2:
    # 이 아래부터는 기존의 [📸 사진 기록 방식] 로직(source_option 등)을 그대로 유지하시면 됩니다.
    st.markdown("### **📸 사진 기록 방식**")
# --- (교체 구간 끝) ---

       
# 1. 업로더 한글화 CSS (디자인 적용)
st.markdown("""
    <style>
        /* 원래 있던 영어 텍스트 숨기기 */
        section[data-testid="stFileUploadDropzone"] div div span,
        section[data-testid="stFileUploadDropzone"] small,
        section[data-testid="stFileUploadDropzone"] button {
            display: none !important;
        }

        /* 상단 안내 문구 */
        section[data-testid="stFileUploadDropzone"] div div::before {
            content: "여기에 사진을 끌어다 놓으세요";
            display: block !important;
            font-size: 0.9rem !important;
            color: #808080 !important;
            margin-bottom: 10px !important;
        }

        /* 빨간색 촬영/선택 버튼 */
        section[data-testid="stFileUploadDropzone"]::before {
            content: "📸 사진 촬영 또는 선택하기";
            display: block !important;
            margin: 10px auto !important;
            padding: 10px 20px !important;
            background-color: #ff4b4b !important;
            color: white !important;
            border-radius: 8px !important;
            cursor: pointer !important;
            font-weight: bold !important;
            text-align: center !important;
            width: fit-content !important;
        }

        /* 하단 용량 제한 문구 */
        section[data-testid="stFileUploadDropzone"] div div::after {
            content: "파일당 최대 200MB • PNG, JPG, JPEG";
            display: block !important;
            font-size: 0.75rem !important;
            color: #a0a0a0 !important;
            margin-top: 5px !important;
        }
    </style>
""", unsafe_allow_html=True)

# 2. 통합된 업로더 (변수명을 img_file로 통일하여 분석 버튼과 연결)
img_file = st.file_uploader(
    "사진 업로드 전용", 
    type=['png', 'jpg', 'jpeg'], 
    label_visibility="collapsed",
    key="safe_upload" # 분석 버튼 로직에서 사용하는 key와 맞춰주는 것이 좋습니다.
)

# 3. 사진 업로드 시 미리보기 및 안내 문구 (요청하신 우측 컬럼 배치)
if img_file:
    col_preview, col_info = st.columns([1, 1.5]) # 왼쪽 미리보기, 오른쪽 문구
    with col_preview:
        st.image(img_file, caption="업로드된 원본 사진", width=300)
    with col_info:
        st.write("") # 상단 여백
        st.success("✅ 사진이 성공적으로 등록되었습니다.")
        st.info("💡 아래 '분석 시작' 버튼을 누르면 AI가 얼굴을 비식별화 후 분석을 시작합니다.")

def apply_face_blur_ai(img_file):
    """
    Gemini AI로 얼굴 좌표를 정밀 탐지하고 OpenCV로 블러링합니다.
    """
    try:
        # 1. 이미지 읽기 및 변환
        img_file.seek(0)
        file_bytes = np.asarray(bytearray(img_file.read()), dtype=np.uint8)
        image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
        if image is None: return img_file.getvalue()
        
        h, w, _ = image.shape
        pil_img = Image.open(io.BytesIO(img_file.getvalue()))

        # 2. Gemini AI에게 얼굴 좌표 요청 (JSON 형식)
        # prompt에 '얼굴이 없다면 빈 리스트를 반환해'라고 명시하여 오류 방지
        prompt = """
        이미지에서 모든 사람의 얼굴(머리 전체) 위치를 찾아서 
        [ymin, xmin, ymax, xmax] 좌표 리스트로 응답해줘. 
        만약 얼굴이 전혀 없다면 {"faces": []} 라고 답해줘.
        JSON 형식: {"faces": [[ymin, xmin, ymax, xmax], ...]}
        """
        
        response = client.models.generate_content(
            model=MODEL_ID,
            contents=[prompt, pil_img],
            config=genai.types.GenerateContentConfig(
                response_mime_type="application/json"
            )
        )

        # 3. 좌표 파싱 및 블러 처리
        face_data = json.loads(response.text)
        faces = face_data.get("faces", [])

        if not faces:
            return img_file.getvalue() # 얼굴이 없으면 원본 반환

        for box in faces:
            ymin, xmin, ymax, xmax = box
            # 상대 좌표를 절대 좌표로 변환 (0~1000 -> 실제 픽셀)
            left, top = int(xmin * w / 1000), int(ymin * h / 1000)
            right, bottom = int(xmax * w / 1000), int(ymax * h / 1000)
            
            rw, rh = right - left, bottom - top
            if rw <= 0 or rh <= 0: continue

            # ROI 추출 및 블러 적용
            face_roi = image[top:bottom, left:right]
            
            # 원형 마스크 생성 (사용자님의 기존 블러 로직 활용)
            mask = np.zeros((rh, rw), dtype=np.uint8)
            cv2.circle(mask, (rw // 2, rh // 2), min(rw, rh) // 2, (255), -1)
            
            # 블러 강도 설정 (level=21 정도가 적당함)
            level = 21
            blurred_roi = cv2.GaussianBlur(face_roi, (level, level), 0)
            
            mask_3ch = cv2.merge([mask, mask, mask])
            image[top:bottom, left:right] = np.where(mask_3ch == 255, blurred_roi, face_roi)

        # 4. 결과 인코딩
        _, buffer = cv2.imencode('.jpg', image, [cv2.IMWRITE_JPEG_QUALITY, 90])
        return buffer.tobytes()

    except Exception as e:
        st.error(f"AI 비식별화 중 오류 발생: {e}")
        return img_file.getvalue()


# --- 6. AI 분석 실행 ---


# --- 분석 시작 버튼 부분 (핵심 로직 통합) ---

if st.button("🚀 {cfg['institution']['abbr']} AI 위험요인 분석 시작", use_container_width=True):
    if not user_description.strip() and not img_file:
        st.warning("⚠️ 분석할 내용(글 또는 사진)을 입력해 주세요.")
    else:
        try:
            with st.spinner(f"✨ KYWA AI가 [{selected_facility}] 시설의 데이터를 분석 중입니다...🔍"):
                # 1. 여기서 무조건 변수를 먼저 만듭니다 (중요!)
                content = []
                
                # 2단계: 분석 프롬프트 구성 (문구 수정 없음)
                prompt = f"""
                {cfg['ai_settings']['persona']}
                
                [시설 정보]
                - 시설명: {selected_facility}
                - 담당부서: {selected_dept}
                - 현장 상황: {user_description}

                [장소 특정 규칙 - 중요]
                1. location(장소) 필드는 [상황 설명 입력]에 적힌 내용 중에서 추출하십시오. 
                   - 예: 사용자가 "본관 2층 테라스 난간 흔들림"이라고 적었다면 -> '본관 3층 복도'
                   - 예: 사용자가 "생활관 계단 논슬립 불량"이라고 적었다면 -> '생활관 계단'
                   - 예: 사용자가 "정문 앞 보도블럭 들뜸"이라고 적었다면 -> '정문 앞'
                3. 만약 상황 설명에 장소 정보가 전혀 없다면, 사진의 시각적 특징을 보고 추론하십시오. (예: '기계실', '화장실', '계단' 등)

                [관련 근거(law) 작성 규칙 - 중요]
                1. 법적 근거는 반드시 기술적 안전 기준이 명시된 현행 법령만 사용하십시오.
                2. 권장 법령: 산업안전보건법, 산업안전보건기준에 관한 규칙, 시특법, 소방기본법, 전기안전관리법 등.
                3. 제외 법령: 청소년활동 진흥법, 일반 행정법 등 실제 안전 기술 기준이 없는 법령은 절대 금지.
                4. 구체적인 조항(예: 제00조)을 언급할 수 있다면 최대한 구체적으로 작성하며, 항이 실제 존재하는지 검증 후 출력하십시오.
                5. 주의 사항: 현행 법령이 아니거나 삭제된 법령 또는 없는 법령에 대해 절대 작성금지.
                  (1) 산업안전보건법 (https://www.law.go.kr/법령/산업안전보건법)
                  (2) 산업안전보건기준에 관한 규칙 (https://www.law.go.kr/법령/산업안전보건기준에관한규칙)
                  (3) 시설물안전법 (https://www.law.go.kr/법령/시설물의안전및유지관리에관한특별법)
                  (4) 소방기본법 (https://www.law.go.kr/법령/소방기본법)
                  (5) 전기안전관리법 (https://www.law.go.kr/법령/전기안전관리법)
                6. 자주 사용되는 조항은 아래 예시를 참고할 것. 다만 조항 번호가 틀리는 것은 안전 관리상 심각한 결함이므로, 확신이 없으면 번호를 생략하십시오.
                  (1) 산업안전보건법 제38조(안전조치)
                  (2) 산업안전보건법 제39조(보건조치)
                  (3) 통로 및 보행 안전: 산업안전보건기준에 관한 규칙 제3조(전도의 방지)
                  (4) 바닥 청결 및 물기 제거: 산업안전보건기준에 관한 규칙 제4조(작업장의 청결)
                  (5) 개인보호구(보안경/안전모 등) 지급: 산업안전보건기준에 관한 규칙 제32조(보호구의 지급 등)
                  (6) 추락 위험 방지: 산업안전보건기준에 관한 규칙 제42조(추락의 방지)
                  (7) 전기 기계/기구 접지: 산업안전보건기준에 관한 규칙 제302조(전기 기계·기구의 접지)
                  (8) 유해물질 취급 시 보호구: 산업안전보건기준에 관한 규칙 제450조(호흡용 보호구의 지급 등)

                [필수 지시사항]
                1. category(분류): 시설명이 '생태', '해양' 등이라 하더라도 이를 category에 적지 마십시오. 
                   반드시 [보행 안전, 시설 안전, 화재 안전, 작업 안전, 활동 안전, 보건 및 위생관리, 화학물질 관리, 작업 환경, 작업 특성, 기계(설비)적 요인, 전기적 요인, 재난 안전] 중 상황에 가장 적합한 표준 분류를 선택하십시오.
                2. 등급 판정의 객관성: 단순 노후화나 경미한 파손(예: 보도블럭 일부 들뜸/파손)은 강도(s)를 2 이하로 설정하여 전체 score가 6점 이하가 되도록 하십시오.

                [시스템 지침] 유해위험요인 분류 가이드라인
                1. 시설 / 화재 / 재난 (구조적·비상 상황)
                 - 시설 안전: 건축물 자체의 노후화, 난간/바닥/천장 등 고정 시설물의 파손 및 구조적 결함. (예: 벽체 균열, 계단 난간 흔들림, 천장재 이탈)
                 - 화재 안전: 화재 발생 가능성 또는 소방 시설의 관리 상태. (예: 소화기 미비, 적치물로 인한 비상구 폐쇄, 가연물 방치)
                 - 재난 안전: 자연재해(강풍, 호우, 지진 등)에 의한 2차 피해 가능성. (예: 옥외 간판 고정 불량, 축대 붕괴 위험, 상습 침수 구역 관리)
                2. 설비 / 전기 / 작업 (기술적·물리적 위험)
                 - 작업 안전: 작업 방식의 위험성 및 안전 장구 착용 상태. (예: 고소 작업 시 안전대 미착용, 사다리 2인 1조 미준수)
                 - 작업 특성: 작업의 강도, 시간, 자세 등 인적 오류 유발 요인. (예: 부적절한 중량물 취급 자세, 무리한 단독 작업, 장시간 반복 동작)
                 - 기계(설비)적 요인: 기계 장치의 기계적 결함 및 안전 장치 유무. (예: 회전부 방호덮개 부재, 승강기 정기 검사 미실시)
                 - 전기적 요인: 전기 설비의 직접적인 위험 및 관리 상태. (예: 피복 노출 전선, 분전함 개방, 문어발식 배선)
                3. 물질 / 환경 / 보건 (건강 및 위생)
                 - 보건 및 위생관리: 위생 상태 및 질병 예방 관리. (예: 화장실/식당 위생 불량, 감염병 방역 미흡)
                 - 화학물질 관리: 유해 화학물질의 저장 및 취급 상태. (예: MSDS 미비치, 유해 물질 용기 라벨 누락, 환기 장치 미가동)
                 - 작업 환경: 작업장의 물리적 환경 요인. (예: 조도 부족으로 인한 어두운 통로, 극심한 소음, 현장 내 분진 발생)
                4. 보행 / 활동 (동적·일상 요인)
                 - 보행 안전: 일상적 이동 경로 상의 위험. (예: 통로 내 돌출물, 바닥 기름/물기로 인한 미끄럼, 통로 조명 불량)
                 - 활동 안전: 야외 활동 및 일반 행위 중 발생하는 위험. (예: 운동 시설 파손, 이동 중 스마트폰 사용 행위, 안전 수칙 미준수 활동)

                [분류 결정 우선순위]
                1. 다중 분류 금지: 가장 지배적인 위험 요인 1가지만 선택하십시오.
                2. 인과관계 고려: 위험의 '원인'이 명확하다면 원인 위주로 분류하십시오. (예: 전기 합선으로 인한 화재 위험은 '전기적 요인'으로 우선 분류)
                3. 보행 vs 시설: 단순히 바닥이 더러운 것은 '보행 안전', 바닥 타일 자체가 깨진 것은 '시설 안전'으로 분류하십시오.
                4. 작업 안전 vs 특성: 안전 보호구 미착용은 '작업 안전', 작업자의 불안정한 신체 자세는 '작업 특성'으로 분류하십시오.

                [빈도 등급 판정 가이드라인] ※1~5번 기준과 예를 근거로 하되 안전수칙 및 작업표준은 있음을 전제로 등급 판정.
                1. 빈도 5점(기준: 피해가 발생할 가능성이 매우 높음(주 1회 발생), 해당 안전대책이 되어 있지 않고, 표시･표지가 있어도 불비(不備)가 많으며, 안전수칙･작업표준 등도 없음, 유해화학물질(증기,분진 등) : 노출수준 법적기준 이상, 직업병 유소견자 발생 시, 근골격계 : 초과근무(1일 8시간 이상))
                2. 빈도 4점(기준: 피해가 발생할 가능성이 높음(1개월에 1회 발생), 가드･방호덮개, 기타 안전장치가 없거나 상당한 불비(不備)가 있고, 비상정지장치, 표시･표지는 웬만큼 설치되어 있으며, 안전수칙･작업표준 등은 있지만 지키기 어렵고 많은 주의를 해야 함, 유해화학물질(증기,분진 등) : 노출수준 법적기준 70%~100%미만, 근골격계 : 계속(1일 4시간 이상))
                3. 빈도 3점(기준: 부주의하면 피해가 발생할 가능성이 있음(1년에 1회 발생), 가드･방호덮개 또는 안전장치 등은 설치되어 있지만, 가드가 낮거나 간격이 벌어져 있는 등 불비(不備)가 있고, 위험영역 접근, 위험원과의 접촉이 있을 수 있으며, 안전수칙･작업표준 등은 있지만 일부 준수하기 어려운 점이 있음, 해화학물질(증기,분진 등) : 노출수준 법적기준 40%~70%미만, 골격계 : 자주(1일 4시간 미만))
                4. 빈도 2점(기준: 피해가 발생할 가능성이 낮음(3년에 1회 발생), 가드･방호덮개 등으로 보호되어 있고, 안전장치가 설치되어 있으며, 위험영역에의 출입이 곤란한 상태이고, 안전수칙･작업표준(서) 등이 정비되어 있고 준수하기 쉬우나, 피해의 가능성이 남아 있음, 유해화학물질(증기,분진 등) : 노출수준 법적기준 10%~40%미만, 골격계 : 가끔(하루 또는 주 2~3일))
                5. 빈도 1점(기준: 피해가 발생할 가능성이 매우 낮음(10년에 1회 발생), 전반적으로 안전조치가 잘 되어 있음, 유해화학물질(증기,분진 등) : 노출수준 법적기준 10%미만, 근골격계 : 3개월 마다(년2~3회))

                [강도 등급 판정 가이드라인]
                1. 강도 4점(사망 또는 영구적 근로불능으로 연결되는 부상․질병(업무에 복귀 불가능), 장애가 남는 부상․질병, 소음 : 법적기준 이상노출(90dB이상), 직업병 발생, 근골격계 : 매우 힘듦 )
                2. 3점(휴업을 수반하는 중대한 부상 또는 질병(일정시점에서는 업무에 복귀 가능(완치 가능)), 소음 : 90dB미만~85dB, 근골격계 : 힘듦)
                3. 2점(응급조치 이상의 치료가 필요하지만 휴업이 수반되지 않는 부상 또는 질병, 소음 : 85dB미만~80dB, 근골격계 : 보통)
                4. 1점(처치(치료) 후 바로 원래의 작업을 수행할 수 있는 경미한 부상 또는 질병(업무에 전혀 지장이 없음))

                [판정 원칙 및 예외 기준]
                1. 일상적 위험 vs 산업적 위험 구분: 단순 전도 등은 강도 1점을 원칙으로 함.
                2. 점수 조정 예시: 보도블럭 파손(빈도 2, 강도 1 -> 2점), 바닥 물기(빈도 3, 강도 1 -> 3점), 키보드 및 마우스 작업(빈도 3, 강도 1 -> 3점)

                [종합 등급 판정 가이드라인]
                - 매우 낮음(1~3점), 낮음(4~6점), 보통(8점), 약간 높음(9~12점), 높음(13~15점), 매우 높음(16~20점)
                - 5점부터는 '허용 불가능한 수준'의 사안으로 판단하므로 경미한 사항은 최대 4점을 기준으로 함.
                - 관리기준([매우낮음]: 안전보건 정보 제공 및 반복적인 안전교육, 실시가 필요한 위험. [낮음]: 안전보건표지 부착, 작업절차서 마련 등 관리적 대책 마련이 필요한 위험. [보통]:  계획된 정비·보수기간에 위험성 감소대책을 세워야 하는 위험. [약간 높음]: 긴급 임시안전대책을 세운 후 작업 또는 운영을 하되 계획된 정비·보수기간에 안전대책을 세워야 하는 위험. [높음]: 즉시 개선을 실행해야 하는 위험. [매우 높음]: 즉시 작업/운영 중단)
                - 모든 문장은 명사형 종결.
                - 반드시 다음 JSON 형식을 엄수하세요: 키는 category, location, scenario, p, s, score, grade, law, solution 이며 리스트 [] 안에 담아 출력하세요.
                """

                # [필수 추가] 생성한 프롬프트를 리스트에 담습니다.
                content.append(prompt)

                # 1단계: 사진이 있다면 즉시 비식별화 처리
                if img_file:
                    # 저희가 위에서 정의한 AI 비식별화 함수 호출
                    processed_bytes = apply_face_blur_ai(img_file)
                    # 비식별화된 바이트 데이터를 Gemini가 읽을 수 있는 PIL 이미지로 변환
                    analysis_image = Image.open(io.BytesIO(processed_bytes))
                    # 사진도 리스트에 담습니다.
                    content.append(analysis_image)

                # [2단계] 재시도 로직 및 최신 라이브러리 호출
                import time
                response = None
                max_retries = 3

                for attempt in range(max_retries):
                    try:
                        # 최신 client.models.generate_content 방식 적용
                        response = client.models.generate_content(
                            model=MODEL_ID,
                            contents=content,
                            config={
                                "response_mime_type": "application/json",
                                "temperature": 0.0
                            }
                        )
                        break  # 성공 시 탈출
                        
                    except Exception as e:
                        error_msg = str(e).lower()
                        # 사용량 초과 및 서버 과부하 에러 시 재시도
                        if any(x in error_msg for x in ["429", "quota", "503", "resource_exhausted"]):
                            if attempt < max_retries - 1:
                                wait_time = 2 * (attempt + 1)
                                time.sleep(wait_time)
                                st.toast(f"⏳ 사용량 조절 중... 재시도 {attempt+1}/{max_retries}")
                                continue
                            else:
                                st.error("🚨 현재 AI 이용량이 많아 분석이 어렵습니다. 잠시 후 다시 시도해주세요.")
                                st.stop()
                        else:
                            st.error(f"❌ 분석 중 오류가 발생했습니다: {e}")
                            st.stop()

                # 3단계: Gemini API 호출 (최신 라이브러리 방식)
                response = client.models.generate_content(
                    model=MODEL_ID,
                    contents=content, # 이제 content가 프롬프트와 사진을 모두 포함합니다.
                    config={
                        "response_mime_type": "application/json",
                        "temperature": 0.0
                    }
                )

                # --- 분석 완료 처리 부분 (기존 코드 하단) ---
                if response:
                    res_data = json.loads(response.text.strip())
                    st.session_state.analysis_results = res_data if isinstance(res_data, list) else [res_data]
    
                    # ✅ 분석에 사용된 비식별 이미지를 전송용으로 세션에 저장
                    if img_file:
                        # analysis_image를 바이트로 변환하여 저장
                        img_byte_arr = io.BytesIO()
                        analysis_image.save(img_byte_arr, format='JPEG')
                        st.session_state.final_secure_image = img_byte_arr.getvalue()
    
                    # 성공 메시지와 리런은 반드시 'if response:' 블록 안에 있어야 합니다.
                    st.success(f"✅ [{selected_facility}] 시설 분석 완료!")
                    st.rerun()

        except Exception as e:
            st.error(f"❌ 최종 처리 중 오류가 발생했습니다: {e}")

# --- 7. 결과 표시 및 데이터 처리 ---
if st.session_state.analysis_results:
    st.markdown("### 📋 AI 위험성평가 결과")
    st.info("💡 **'위험상황'**과 **'감소대책'** 칸을 클릭하여 직접 내용을 수정할 수 있습니다.")

    # 1. 데이터를 데이터프레임으로 변환
    df = pd.DataFrame(st.session_state.analysis_results)

    # 2. 데이터 에디터 설정
    edited_df = st.data_editor(
        df,
        column_config={
            "category": st.column_config.TextColumn("분류", disabled=True),
            "location": st.column_config.TextColumn("📍장소(편집 가능)", width="midium"), # 추가된 부분
            "scenario": st.column_config.TextColumn(
                "✅ 위험상황 (편집 가능)", 
                help="현장 상황에 맞춰 내용을 수정하세요.",
                width="medium"
            ),
            "p": st.column_config.TextColumn("빈도", disabled=True, width="small"),
            "s": st.column_config.TextColumn("강도", disabled=True, width="small"),
            "score": st.column_config.TextColumn("점수", disabled=True, width="small"),
            "grade": st.column_config.TextColumn("등급", disabled=True, width="small"),
            "law": st.column_config.TextColumn("관련근거", disabled=True, width="medium"),
            "solution": st.column_config.TextColumn(
                "✅ 감소대책 (편집 가능)", 
                help="현장에 맞는 대책으로 수정하세요.",
                width="large",
                required=True
            )
        },
        disabled=["category", "p", "s", "score", "grade", "law"],
        width="stretch",
        hide_index=True,
        key="final_editor_main"
    )

    # 수정된 데이터를 즉시 세션 상태에 업데이트
    st.session_state.final_data = edited_df.to_dict('records')

    # --- [3단계] 전송 버튼 로직 ---
    st.write("")
    if st.button("✅ {cfg['institution']['abbr']} 안전센터로 데이터 최종 전송", use_container_width=True):
        if sheets_service is None or drive_service is None:
            st.error("⚠️ GCP 인증에 실패하여 데이터를 전송할 수 없습니다. 관리자에게 문의하세요.")
        elif not st.session_state.get("final_data"):
            st.error("⚠️ 전송할 데이터가 없습니다.")
        else:
            with st.spinner("🚀 KYWA AI 안전센터로 데이터를 전송 중입니다..."):
                try:
                    # 한국 시간 설정
                    now_kst = datetime.datetime.now() + datetime.timedelta(hours=9)
                    current_time = now_kst.strftime("%Y-%m-%d %H:%M:%S")
                    timestamp_str = now_kst.strftime("%Y%m%d_%H%M%S")

                    photo_link = "사진 없음"
                    
                    # 비식별 이미지 확인 및 드라이브 업로드
                    if "final_secure_image" in st.session_state and st.session_state.final_secure_image:
                        filename = f"{selected_facility}_{timestamp_str}.jpg"
                        # 앱스 스크립트 대신, 1단계에서 만든 서비스 계정 업로드 함수를 직접 호출합니다.
                        photo_link = upload_to_drive(
                            file_name=filename, 
                            file_content=st.session_state.final_secure_image, 
                            mime_type='image/jpeg'
                        )
                    
                    success_count = 0

                    # 구글 시트 컬럼 순서에 맞춰 리스트 재구성
                    for row in st.session_state.final_data:
                        sheet_row = [
                            current_time,           # 타임스탬프
                            selected_facility,      # 시설명
                            selected_dept,          # 담당 부서
                            row.get("location"),    # 장소
                            row.get("category"),    # 유해위험요인(분류)
                            row.get("scenario"),    # 위험상황
                            row.get("p"),           # 빈도
                            row.get("s"),           # 강도
                            row.get("score"),       # 점수
                            row.get("grade"),       # 위험등급
                            row.get("solution"),    # 감소대책
                            row.get("law"),         # 관련근거
                            photo_link              # 사진 기록
                        ]
                        
                        if append_row_to_sheet(sheet_row):
                            success_count += 1
                    
                    if success_count > 0:
                        st.success(f"✅ 데이터 {success_count}건이 성공적으로 전송되었습니다!")
                        st.balloons()
                        # 전송 후 데이터 초기화 (필요 시 주석 해제)
                        # st.session_state.analysis_results = None
                        
                except Exception as e:
                    st.error(f"❌ 전송 중 오류 발생: {e}")


    # --- 저장 버튼 영역 (분석 직후 바로 나타나며, 클릭 시 사라짐 방지) ---
    st.markdown("---")
    dl_col1, dl_col2 = st.columns(2)
    
    # 수정된 데이터를 실시간으로 함수에 전달
    if st.session_state.final_data:
        with dl_col1:
            st.download_button(
                label="📂 Word 저장",
                data=create_docx(st.session_state.final_data),
                file_name=f"KYWA_{selected_facility}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                width="stretch",
                key="btn_word_download" # 클릭 시 사라짐 방지를 위한 고유 키
            )
        with dl_col2:
            st.download_button(
                label="📊 Excel 저장",
                data=create_excel(st.session_state.final_data),
                file_name=f"KYWA_{selected_facility}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                width="stretch",
                key="btn_excel_download" # 클릭 시 사라짐 방지를 위한 고유 키
            )

# --- [수정] 날짜 형식 오류를 해결한 데이터 로드 함수 ---
def load_dashboard_data():
    # config.json에 정의된 SPREADSHEET_ID를 동적으로 불러옴 (첫 번째 시트 기준)
    sheet_url = f"https://docs.google.com/spreadsheets/d/{SPREADSHEET_ID}/export?format=csv"
    
    try:
        df = pd.read_csv(sheet_url)
        
        if '타임스탬프' in df.columns:
            # 1. 한국어 '오전/오후'를 Pandas가 인식 가능한 'AM/PM'으로 변경
            df['타임스탬프'] = df['타임스탬프'].str.replace('오전', 'AM').str.replace('오후', 'PM')
            
            # 2. 날짜 형식으로 변환 (format을 지정하지 않아도 치환 후에는 잘 작동합니다)
            df['타임스탬프'] = pd.to_datetime(df['타임스탬프'], format='mixed', errors='coerce')
            
            # 3. 변환 실패한 데이터(NaT) 제거 (선택 사항)
            df = df.dropna(subset=['타임스탬프'])
            
        return df
    except Exception as e:
        st.error(f"데이터를 불러오는 중 오류가 발생했습니다: {e}")
        return None

# --- 대시보드 섹션 ---
st.write("---")
dashboard_data = load_dashboard_data()

if dashboard_data is not None:
    # 2. 날짜 필터링 (2026년 데이터만)
    if '타임스탬프' in dashboard_data.columns:
        yearly_data = dashboard_data[dashboard_data['타임스탬프'].dt.year == 2026].copy()
    else:
        yearly_data = dashboard_data.copy()

    if yearly_data.empty:
        st.warning("📅 2026년도 데이터가 아직 없습니다. 데이터를 첫 번째로 전송해 보세요!")
    else:
        st.subheader("📊 실시간 점검 데이터 현황 (2026년)")
        
        # 3. 상단 지표
        total_count = len(yearly_data)
        m1, m2 = st.columns(2)
        m1.metric("올해 누적 점검 건수", f"{total_count} 건")
        
        author_col = "작성자 성명" 
        if author_col in yearly_data.columns:
            m2.metric("참여 인원(명)", f"{yearly_data[author_col].nunique()} 명")
        else:
            m2.metric("점검결과 제출 시설", f"{yearly_data['시설명'].nunique()} 개 시설")

        # --- 색상 맵 설정 ---
        CATEGORY_COLOR_MAP = {
        # 1. 시설/재난/화재 (고정 위험 및 비상사태 - 붉은색 계열)
            "시설 안전": "#D32F2F",      # 진한 빨강
            "화재 안전": "#FF5722",      # 주황빛 빨강 (불꽃)
            "재난 안전": "#880E4F",      # 자주색 (중대 재난)

        # 2. 설비/전기/작업 (기술적/물리적 요인 - 노란색/갈색 계열)
            "작업 안전": "#FFA000",      # 호박색 (주의/작업)
            "작업 특성": "#E64A19",      # 진한 주황 (인적 요인/작업 강도)
            "기계(설비)적 요인": "#795548", # 갈색 (기계/금속)
            "전기적 요인": "#FBC02D",    # 노란색 (전기/번개)

        # 3. 물질/환경/보건 (보이지 않는 위해 요인 - 보라색/회색 계열)
            "보건 및 위생관리": "#E91E63", # 분홍/보라 (의료/위생)
            "화학물질 관리": "#9C27B0",   # 보라 (유독물질)
            "작업 환경": "#455A64",      # 블루그레이 (환경/소음/먼지)

        # 4. 보행/활동 (동적 유해요인 - 파란색/초록색 계열)
            "보행 안전": "#1976D2",      # 파란색 (통로/이동)
            "활동 안전": "#388E3C"       # 초록색 (일상 활동/야외)
        }

        FACILITY_COLOR_MAP = {
            "중앙": "#B93444", "본원": "#6B5B95", "평창": "#E2725B",
            "바이오": "#D2B48C", "해양": "#5B84B1", "우주": "#2E4A62",
            "미래": "#92B06A", "생태": "#5F7161"
        }

# --- 4. 그래프 시각화 영역 (3단 구성으로 변경) ---
        g_col1, g_col2, g_col3 = st.columns(3)

        # [1단] 장소 현황 (D컬럼 - Index 3)
        with g_col1:
            if len(yearly_data.columns) >= 4:
                target_col_loc = yearly_data.columns[3] 
                st.write(f"**📍 {target_col_loc} 현황**")
                if not yearly_data[target_col_loc].dropna().empty:
                    yearly_data[target_col_loc] = yearly_data[target_col_loc].astype(str).str.strip()
                    
                    fig_loc = px.pie(
                        yearly_data, names=target_col_loc, hole=0.3,
                        color_discrete_sequence=px.colors.qualitative.Pastel
                    )
                    fig_loc.update_traces(
                        textinfo='percent+value', 
                        texttemplate='%{percent:.0%}<br>(%{value}건)',
                        insidetextorientation='horizontal',
                        textfont_size=11
                    )
                    # 🔴 [수정됨] fig_pie -> fig_loc 으로 변경
                    fig_loc.update_layout(
                        margin=dict(t=30, b=80, l=0, r=0), 
                        height=450, 
                        showlegend=True,
                        legend=dict(
                            orientation="h",      
                            yanchor="top",        
                            y=-0.1,               
                            xanchor="center",     
                            x=0.5,
                            font=dict(size=10),   
                            itemwidth=30          
                        ),
                        paper_bgcolor='rgba(0,0,0,0)',
                        dragmode=False
                    )
                    st.plotly_chart(fig_loc, use_container_width=True, config={'displayModeBar': False})

        # [2단] 유해위험요인 현황 (E컬럼 - Index 4)
        with g_col2:
            if len(yearly_data.columns) >= 5:
                target_col_cat = yearly_data.columns[4] 
                st.write(f"**⚠️ {target_col_cat} 현황**")
                if not yearly_data[target_col_cat].dropna().empty:
                    yearly_data[target_col_cat] = yearly_data[target_col_cat].astype(str).str.strip()
                    
                    fig_pie = px.pie(
                        yearly_data, names=target_col_cat, hole=0.3,
                        color=target_col_cat, color_discrete_map=CATEGORY_COLOR_MAP
                    )
                    fig_pie.update_traces(
                        textinfo='percent+value', 
                        texttemplate='%{percent:.0%}<br>(%{value}건)',
                        insidetextorientation='horizontal',
                        textfont_size=11
                    )
                    fig_pie.update_layout(
                        margin=dict(t=30, b=80, l=0, r=0), 
                        height=450, 
                        showlegend=True,
                        legend=dict(
                            orientation="h",      
                            yanchor="top",        
                            y=-0.1,               
                            xanchor="center",     
                            x=0.5,
                            font=dict(size=10),   
                            itemwidth=30          
                        ),
                        paper_bgcolor='rgba(0,0,0,0)',
                        dragmode=False
                    )
                    st.plotly_chart(fig_pie, use_container_width=True, config={'displayModeBar': False})

        # [3단] 시설별 점검 건수
        with g_col3:
            target_col_fac = "시설명" 
            if target_col_fac in yearly_data.columns:
                st.write(f"**🏢 {target_col_fac}별 건수**")
                yearly_data[target_col_fac] = yearly_data[target_col_fac].astype(str).str.strip()
                fac_counts = yearly_data[target_col_fac].value_counts().reset_index()
                fac_counts.columns = [target_col_fac, '건수']
                
                fig_bar = px.bar(
                    fac_counts, x=target_col_fac, y='건수', color=target_col_fac,
                    color_discrete_map=FACILITY_COLOR_MAP
                )
                fig_bar.update_traces(
                    texttemplate='%{y}건', 
                    textposition='outside',
                    textfont_size=11
                )
                fig_bar.update_layout(
                    margin=dict(t=35, b=0, l=0, r=0), height=450, # 높이를 1, 2단과 동일하게 450으로 맞춤 
                    showlegend=False,
                    xaxis_title=None, yaxis_title=None,
                    paper_bgcolor='rgba(0,0,0,0)',
                    plot_bgcolor='rgba(0,0,0,0)',
                    dragmode=False 
                )
                st.plotly_chart(fig_bar, use_container_width=True, config={'displayModeBar': False})

# --- [5] 동적 푸터(Footer) 섹션 (교체 구간) ---
st.write("") # 간격 확보
st.write("---")
footer_cols = st.columns([3, 1])

with footer_cols[0]:
    st.markdown(f"### 🔒 Data Governance & Privacy")
    st.caption(f"""
    **© 2026 {cfg['institution']['name']} {cfg['footer']['department']}.** 본 시스템은 **공공기관 AI 활용 가이드라인** 및 **정보보안 업무규정**을 엄격히 준수합니다.
    
    * **데이터 보안:** 입력된 모든 정보는 **API 옵트아웃(Opt-out) 설정**이 적용되어 외부 모델 학습에 활용되지 않습니다.
    * **운영 방침:** **{cfg['institution']['abbr']} AI 안전센터**로 전송된 데이터는 **담당자의 데이터 정합성 검토**를 거칩니다. 
      점검 내용이 부적절하거나 중복된 경우, 데이터 신뢰성 유지를 위해 운영 관리자에 의해 임의 수정 또는 삭제될 수 있습니다.
    * **면책 고지:** AI 분석 정보는 위험 요인 발굴을 돕는 가이드라인입니다. 실제 위험성 평가 시에는 현장 상황을 반영한 담당 직원의 면밀한 검토를 권고합니다.
    """)

with footer_cols[1]:
    st.markdown("### 📞 Contact")
    # HTML을 사용하여 config 정보를 동적으로 삽입합니다.
    st.markdown(f"""
    <div style="line-height: 1.6;">
        <span style="font-weight: bold; font-size: 0.9rem; color: #31333F;">{cfg['footer']['department']}</span><br>
        <span style="color: #444; font-size: 0.85rem;">📧 {cfg['footer']['email']}</span><br>
        <span style="color: #444; font-size: 0.85rem;">
            <span style="display: inline-block; transform: rotate(10deg); color: #000;">📞</span> {cfg['footer']['phone']}
        </span>
    </div>
    """, unsafe_allow_html=True)

# 최하단 한 줄 강조 (기관 약어와 앱 타이틀 동적 반영)
st.markdown(f"""
    <p style='font-size: 0.8rem; color: gray; text-align: center;'>
        Safe Together, {cfg['institution']['abbr']} {cfg['institution']['app_title']}
    </p>
""", unsafe_allow_html=True)
